from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "简益咨询-产品与服务白皮书.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)
ILLUSTRATIONS = ROOT / "assets" / "whitepaper" / "illustrations"

# standard_business_brief + named overrides:
# a4_china_business_document, zh_business_typography, jianyi_brand_green
PAGE_W_DXA = 11906
MARGIN_DXA = 1440
CONTENT_DXA = PAGE_W_DXA - 2 * MARGIN_DXA
TABLE_INDENT_DXA = 130
CELL_TOP_BOTTOM = 90
CELL_START_END = 130

GREEN = "1F4E3D"
GREEN_2 = "355F50"
BLACK = "1A1A1A"
DARK_GRAY = "4A4A4A"
MID_GRAY = "6B6B6B"
LIGHT_GRAY = "F2F4F3"
LINE_GRAY = "C8CECB"
WHITE = "FFFFFF"

BODY_CN = "Heiti SC"
HEAD_CN = "Heiti SC"
LATIN = "Arial"


def set_rfonts(run, *, cn=BODY_CN, latin=LATIN, size=None, bold=None,
               color=None, italic=None):
    # LibreOffice does not consistently honor Word's eastAsia fallback when
    # ascii/hAnsi point to a Latin-only font. Use the Chinese family as the
    # primary family so both Word and the PDF export retain all glyphs.
    run.font.name = cn
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), cn)
    rfonts.set(qn("w:hAnsi"), cn)
    rfonts.set(qn("w:eastAsia"), cn)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, *, cn, latin, size, bold=False, color=BLACK):
    style.font.name = cn
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), cn)
    rfonts.set(qn("w:hAnsi"), cn)
    rfonts.set(qn("w:eastAsia"), cn)


def set_cell_margins(cell, top=CELL_TOP_BOTTOM, start=CELL_START_END,
                     bottom=CELL_TOP_BOTTOM, end=CELL_START_END):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = tc_mar.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_mar.append(tag)
        tag.set(qn("w:w"), str(value))
        tag.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=LINE_GRAY, size=5):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_DXA, (sum(widths), CONTENT_DXA)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, *, side="left", color=GREEN, size=18, space=8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_rfonts(run, cn=HEAD_CN, size=8.5, color=MID_GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, sep, text, end):
        run._r.append(element)


def add_numbering_definition(doc, num_fmt="bullet"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl.append(fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if num_fmt == "bullet" else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "100")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    if num_fmt == "bullet":
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), LATIN)
        r_fonts.set(qn("w:hAnsi"), LATIN)
        r_pr.append(r_fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_list(doc, items, *, numbered=False):
    num_id = add_numbering_definition(doc, "decimal" if numbered else "bullet")
    for item in items:
        p = doc.add_paragraph(style="List Body")
        apply_num(p, num_id)
        run = p.add_run(item)
        set_rfonts(run, size=10.5, color=BLACK)
    return num_id


def add_body(doc, text, *, bold_lead=None, keep=False):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_rfonts(r1, cn=HEAD_CN, size=10.5, bold=True, color=BLACK)
        r2 = p.add_run(text[len(bold_lead):])
        set_rfonts(r2, size=10.5, color=BLACK)
    else:
        run = p.add_run(text)
        set_rfonts(run, size=10.5, color=BLACK)
    p.paragraph_format.keep_together = keep
    return p


def add_callout(doc, label, text):
    p = doc.add_paragraph(style="Callout")
    set_paragraph_shading(p, LIGHT_GRAY)
    set_paragraph_border(p, side="left", color=GREEN, size=20, space=8)
    label_run = p.add_run(f"{label}  ")
    set_rfonts(label_run, cn=HEAD_CN, size=10.5, bold=True, color=GREEN)
    body_run = p.add_run(text)
    set_rfonts(body_run, size=10.5, color=BLACK)
    return p


def add_kv(doc, label, text):
    p = doc.add_paragraph(style="Key Value")
    r1 = p.add_run(f"{label}：")
    set_rfonts(r1, cn=HEAD_CN, size=10.5, bold=True, color=GREEN)
    r2 = p.add_run(text)
    set_rfonts(r2, size=10.5, color=BLACK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_section_illustration(doc, filename, caption, alt_text):
    image_path = ILLUSTRATIONS / filename
    if not image_path.exists():
        raise FileNotFoundError(image_path)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(image_path), width=Mm(159))
    doc_pr = shape._inline.docPr
    doc_pr.set("title", caption)
    doc_pr.set("descr", alt_text)

    cp = doc.add_paragraph(style="Image Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.keep_together = True
    cr = cp.add_run(caption)
    set_rfonts(cr, size=8.5, color=MID_GRAY)
    return shape


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(text)
        set_rfonts(run, cn=HEAD_CN, size=9, bold=True, color=WHITE)
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            set_cell_margins(cell)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.18
            run = p.add_run(text)
            set_rfonts(run, cn=HEAD_CN if idx == 0 else BODY_CN,
                       size=font_size, bold=(idx == 0), color=BLACK)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, cn=BODY_CN, latin=LATIN, size=10.5, color=BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.30

    for name, size, color, before, after in (
        ("Heading 1", 16, GREEN, 16, 8),
        ("Heading 2", 13, GREEN, 12, 6),
        ("Heading 3", 11.5, DARK_GRAY, 8, 4),
    ):
        st = styles[name]
        set_style_font(st, cn=HEAD_CN, latin=LATIN, size=size, bold=True, color=color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.10
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.keep_together = True

    title = styles["Title"]
    set_style_font(title, cn=HEAD_CN, latin=LATIN, size=28, bold=True, color=GREEN)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.line_spacing = 1.05
    title_ppr = title.element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, cn=HEAD_CN, latin=LATIN, size=14, color=DARK_GRAY)
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle_ppr = subtitle.element.get_or_add_pPr()
    subtitle_border = subtitle_ppr.find(qn("w:pBdr"))
    if subtitle_border is not None:
        subtitle_ppr.remove(subtitle_border)

    for style_name in ("List Body", "Callout", "Key Value", "Table Citation", "Image Caption"):
        if style_name not in [s.name for s in styles]:
            styles.add_style(style_name, 1)
    list_body = styles["List Body"]
    set_style_font(list_body, cn=BODY_CN, latin=LATIN, size=10.5, color=BLACK)
    list_body.paragraph_format.space_after = Pt(5)
    list_body.paragraph_format.line_spacing = 1.167
    callout = styles["Callout"]
    set_style_font(callout, cn=BODY_CN, latin=LATIN, size=10.5, color=BLACK)
    callout.paragraph_format.left_indent = Mm(4)
    callout.paragraph_format.right_indent = Mm(4)
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.line_spacing = 1.25
    key_value = styles["Key Value"]
    set_style_font(key_value, cn=BODY_CN, latin=LATIN, size=10.5, color=BLACK)
    key_value.paragraph_format.space_after = Pt(5)
    key_value.paragraph_format.line_spacing = 1.25
    citation = styles["Table Citation"]
    set_style_font(citation, cn=BODY_CN, latin=LATIN, size=8.5, color=MID_GRAY)
    citation.paragraph_format.space_before = Pt(4)
    citation.paragraph_format.space_after = Pt(4)
    image_caption = styles["Image Caption"]
    set_style_font(image_caption, cn=BODY_CN, latin=LATIN, size=8.5, color=MID_GRAY)
    image_caption.paragraph_format.space_before = Pt(0)
    image_caption.paragraph_format.space_after = Pt(8)
    image_caption.paragraph_format.line_spacing = 1.05


def set_page_geometry(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(25.4)
    section.header_distance = Mm(12.5)
    section.footer_distance = Mm(12.5)


def configure_section(doc):
    section = doc.sections[0]
    set_page_geometry(section)
    section.different_first_page_header_footer = False

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("简益咨询  JIANYI CONSULTING")
    set_rfonts(r, cn=HEAD_CN, size=8.5, color=MID_GRAY)


def start_body_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_geometry(section)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Mm(159), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run("简益咨询 | 产品与服务白皮书")
    set_rfonts(r1, cn=HEAD_CN, size=8.5, bold=True, color=GREEN)
    r2 = p.add_run("\t零售咨询 · IT规划 · 项目管理 · 产品设计")
    set_rfonts(r2, cn=HEAD_CN, size=8, color=MID_GRAY)
    set_paragraph_border(p, side="bottom", color=LINE_GRAY, size=4, space=5)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(4)
    set_paragraph_border(fp, side="top", color=LINE_GRAY, size=4, space=5)
    r = fp.add_run("简益咨询  |  ")
    set_rfonts(r, cn=HEAD_CN, size=8.5, color=MID_GRAY)
    add_page_field(fp)


def add_cover(doc):
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(10)
    r = kicker.add_run("JIANYI CONSULTING")
    set_rfonts(r, cn=HEAD_CN, size=10, bold=True, color=GREEN)
    title = doc.add_paragraph("简益咨询", style="Title")
    title.paragraph_format.space_after = Pt(2)
    subtitle = doc.add_paragraph("产品与服务白皮书", style="Subtitle")
    subtitle.paragraph_format.space_after = Pt(14)
    set_paragraph_border(subtitle, side="bottom", color=GREEN, size=12, space=12)

    lead = doc.add_paragraph()
    lead.paragraph_format.space_before = Pt(16)
    lead.paragraph_format.space_after = Pt(20)
    lead.paragraph_format.line_spacing = 1.35
    r = lead.add_run("让经营判断、系统方案与项目执行，站在同一套逻辑上。")
    set_rfonts(r, cn=HEAD_CN, size=17, bold=False, color=BLACK)

    service = doc.add_paragraph()
    service.paragraph_format.space_after = Pt(18)
    r = service.add_run("零售经营咨询  |  IT战略与规划  |  甲方项目管理  |  业务产品设计")
    set_rfonts(r, cn=HEAD_CN, size=10.5, bold=True, color=GREEN)

    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(4)
    r = meta.add_run("版本：V1.3  |  日期：2026年8月  |  用途：对外业务交流")
    set_rfonts(r, cn=HEAD_CN, size=9.5, color=MID_GRAY)
    note = doc.add_paragraph()
    r = note.add_run("本文件所列价格为参考起价，实际范围与报价以双方确认的项目建议书为准。")
    set_rfonts(r, size=9, color=MID_GRAY)


def add_guide(doc):
    add_heading(doc, "内容导览", 1)
    add_body(doc, "本白皮书用于帮助零售企业管理层、业务负责人、IT负责人和项目团队快速了解简益咨询能够解决的问题、合作方式、项目方法和代表性实践。")
    items = [
        "核心观点：为什么许多经营与数字化项目难以形成闭环",
        "服务体系：四类服务对应的管理问题与交付结果",
        "产品化合作方案：从诊断、方案设计到实施陪跑",
        "专项服务说明：服务边界、工作内容与主要交付",
        "案例与最佳实践：ERP、品类管理、IT规划与产品设计",
        "项目方法与治理机制：五步工作法及双方责任",
        "如何启动合作：项目讨论前的准备清单",
        "关于简益：跨零售企业、软件公司和新零售平台的经验视角",
    ]
    add_list(doc, items, numbered=True)
    add_callout(doc, "阅读建议", "如果课题尚不清晰，先阅读“核心观点”和“产品化合作方案”；如果项目已经启动，重点阅读“案例与最佳实践”及“项目方法与治理机制”。")
    add_heading(doc, "文档定位", 2)
    add_body(doc, "本文件是服务介绍与项目讨论材料，不替代针对具体企业的诊断、需求分析和正式项目建议书。案例采用匿名化及概括性表达，避免将阶段性项目成果表述为未经验证的长期经营结果。")
    doc.add_page_break()


def add_core(doc):
    add_heading(doc, "1. 核心观点：企业需要的不只是一个方案", 1)
    add_section_illustration(
        doc, "01-core-idea.png",
        "章节导读｜经营目标、业务动作、数据证据与数字系统形成管理闭环",
        "零售现场、管理会议、数据分析与系统架构连接成闭环。",
    )
    add_body(doc, "许多零售项目的问题，并非团队缺少努力或供应商缺少功能，而是经营目标、业务规则、系统能力和项目责任没有在同一框架中被持续管理。")
    add_heading(doc, "1.1 常见表面现象", 2)
    add_list(doc, [
        "报表数量持续增加，但管理动作没有同步形成。",
        "系统逐步增多，但关键流程仍依赖人工协调和线下判断。",
        "项目会议频繁召开，但问题、责任与决策长期悬而未决。",
    ])
    add_heading(doc, "1.2 深层原因与管理后果", 2)
    add_table(doc,
              ["深层原因", "典型表现", "管理后果"],
              [
                  ("目标口径没有统一", "业务、财务和IT使用不同定义讨论同一问题", "决策周期延长，结果难以比较"),
                  ("规则与系统彼此脱节", "业务规则未进入流程和系统，或系统逻辑脱离现场", "返工增加，数据可信度下降"),
                  ("责任边界缺少闭环", "问题没有明确责任人、决策人和验收标准", "风险持续积累，经验难以复制"),
              ],
              [2100, 3600, CONTENT_DXA - 5700], font_size=9.3)
    add_callout(doc, "简益的工作", "把经营问题翻译成可执行的项目，把项目方案组织成可管理的过程，再通过测试、复盘和验收验证交付是否真正发生。")


def add_service_system(doc):
    add_heading(doc, "2. 服务体系", 1)
    add_section_illustration(
        doc, "02-service-system.png",
        "章节导读｜四类专业能力围绕同一个经营问题协同工作",
        "经营咨询、项目管理、产品设计和IT规划围绕零售经营核心协同。",
    )
    add_body(doc, "简益咨询围绕零售经营、IT规划、项目实施和业务产品设计提供四类服务。服务可以单独采购，也可以根据项目阶段组合。")
    add_table(doc,
              ["服务方向", "管理层要回答", "关键工作", "主要交付"],
              [
                  ("S01 零售经营与品类咨询", "经营问题在哪里，应该先改什么？", "经营诊断、指标口径、品类策略、动作设计、复盘机制", "问题地图、行动清单、试点与复盘方案"),
                  ("S02 IT战略、规划与系统选型", "系统为什么建设，先后顺序是什么？", "业务蓝图、应用架构、系统边界、需求框架、选型支持", "IT规划蓝图、项目路线图、选型与决策依据"),
                  ("S03 ERP/WMS实施与甲方PMO", "如何控制范围、风险、进度和验收？", "范围管理、计划推进、跨方协调、测试验收、上线陪跑", "项目台账、风险与决策机制、验收与移交体系"),
                  ("S04 业务产品与数据工具设计", "复杂业务如何转化为可用产品？", "角色流程、状态权限、PRD、数据模型、原型与场景验证", "产品方案、关键原型、可开发需求包"),
              ],
              [1900, 2050, 2650, CONTENT_DXA - 6600], font_size=8.8)
    p = doc.add_paragraph(style="Table Citation")
    r = p.add_run("组合建议：课题尚不清晰，优先采用 S01；系统建设前采用 S02；项目推进复杂时引入 S03；新工具或业务平台建设采用 S04。")
    set_rfonts(r, size=8.5, color=MID_GRAY)
    doc.add_page_break()


def add_packages(doc):
    add_heading(doc, "3. 产品化合作方案", 1)
    add_section_illustration(
        doc, "03-engagement-packages.png",
        "章节导读｜从问题诊断、方案设计到实施陪跑的递进式合作",
        "三个递进阶段依次呈现诊断、方案设计和实施支持。",
    )
    add_body(doc, "合作方案按照客户所处阶段划分。参考起价用于建立预算预期，正式报价将根据项目范围、组织复杂度、现场投入、数据条件和交付深度确定。")
    add_table(doc,
              ["合作方案", "适用情形", "核心范围", "参考起价"],
              [
                  ("经营诊断与项目定义", "问题复杂，方向尚未统一", "关键访谈与资料诊断；问题地图与优先级；项目定义与行动建议", "¥28,000 起"),
                  ("专项咨询与方案设计", "课题明确，需要形成系统方案", "现状与目标模型；业务或系统方案设计；试点、指标与实施路径", "¥88,000 起"),
                  ("实施陪跑与项目管理", "多部门、多供应商或关键上线项目", "计划、范围与风险管理；跨方协同与决策机制；测试验收与上线陪跑", "¥168,000 起"),
              ],
              [2100, 2150, 3350, CONTENT_DXA - 7600], font_size=9.1)
    add_heading(doc, "3.1 报价通常受哪些因素影响", 2)
    add_list(doc, [
        "项目覆盖的业态、区域、门店、部门和系统数量。",
        "是否需要现场驻场、跨城市访谈、供应商协调或上线陪跑。",
        "数据质量、现有文档完整度以及客户团队可投入的资源。",
        "交付物深度，包括诊断报告、业务蓝图、详细方案、PRD、测试与验收材料等。",
    ])
    add_callout(doc, "报价原则", "先确定要解决的问题、工作边界和验收标准，再讨论投入与价格。价格不应脱离项目责任和交付深度单独比较。")


def add_service_details(doc):
    add_heading(doc, "4. 专项服务说明", 1)
    add_section_illustration(
        doc, "04-specialist-services.png",
        "章节导读｜把流程、角色、产品规则和系统规划组织成可执行方案",
        "工作台连接流程、角色、产品蓝图和IT架构四类专业分析。",
    )
    services = [
        ("4.1 零售经营与品类咨询",
         "适用于经营问题缺少统一判断、指标很多但动作不清晰、品类管理方法难以在门店落地的情况。",
         ["经营现状访谈与问题诊断", "指标口径、经营分析和问题优先级", "品类角色、商品结构与门店动作设计", "试点范围、复盘机制和扩大条件设计"],
         ["经营问题地图", "指标与分析框架", "行动清单及责任建议", "试点与复盘方案"]),
        ("4.2 IT战略、规划与系统选型",
         "适用于系统较多但边界不清、建设顺序缺少依据，或新项目需要从业务蓝图出发规划系统能力的情况。",
         ["业务战略与关键能力梳理", "应用架构、数据关系和系统边界设计", "建设优先级、项目分期与组织协同建议", "需求框架、供应商选型和方案评审支持"],
         ["业务与IT规划蓝图", "系统边界及项目路线图", "需求与选型框架", "管理层决策材料"]),
        ("4.3 ERP/WMS实施与甲方PMO",
         "适用于实施范围复杂、跨部门和跨供应商协作频繁，或甲方需要增强项目控制、测试验收与上线管理能力的情况。",
         ["项目范围、里程碑和计划管理", "问题、风险、变更和决策台账", "业务、IT与供应商之间的跨方协调", "测试、上线、验收和移交组织"],
         ["项目治理机制与台账", "关键会议及决策材料", "测试与验收体系", "上线及移交建议"]),
        ("4.4 业务产品与数据工具设计",
         "适用于业务复杂、角色众多、状态和权限难以说清，或需要把管理方法转化为系统工具的情况。",
         ["用户角色、业务流程、状态和权限梳理", "业务规则、异常场景和关键数据模型", "PRD、交互原型及产品方案", "场景化评审、测试与验收标准"],
         ["产品需求文档", "关键流程与交互原型", "业务规则和数据模型", "可开发、可测试的需求包"]),
    ]
    add_body(doc, "四项专项服务分别对应经营改进、系统规划、项目交付和产品设计问题。下表保留服务边界、典型工作和主要交付，便于管理层直接比较与组合采购。")
    rows = []
    for title, intro, works, outputs in services:
        service_name = title.split(" ", 1)[1]
        rows.append((
            service_name,
            intro,
            "；".join(works),
            "；".join(outputs),
        ))
    add_table(
        doc,
        ["专项服务", "适用情形", "典型工作", "主要交付"],
        rows,
        [1600, 2400, 2950, CONTENT_DXA - 6950],
        font_size=8.4,
    )


def add_cases(doc):
    doc.add_page_break()
    add_heading(doc, "5. 案例与最佳实践", 1)
    add_section_illustration(
        doc, "05-cases.png",
        "章节导读｜用项目路径和形成的管理基础说明实践价值",
        "ERP、品类管理、购物中心IT规划和业务产品设计四类匿名案例。",
    )
    add_body(doc, "以下案例来自过往项目经验的概括性表达。内容重点说明问题结构、工作方法和形成的管理基础，不以未经持续验证的经营数字作为宣传结论。")

    add_heading(doc, "5.1 ERP与财务管理：从系统实施走向管理闭环", 2)
    add_kv(doc, "项目命题", "如何让ERP同时服务业务运营、财务核算与管理分析？")
    add_kv(doc, "简益角色", "业务与财务口径梳理、流程及系统方案协同、项目推进与问题闭环、测试验证与管理应用衔接。")
    add_heading(doc, "工作路径", 3)
    add_list(doc, [
        "统一口径：连接业务数据与财务系统，明确核算和管理分析基础。",
        "组织流程：将应收、返利、费用与经营责任放入可追踪流程。",
        "系统验证：围绕关键场景测试规则、数据和操作结果。",
        "管理应用：支持按部门、柜组等维度查看经营与利润表现。",
    ], numbered=True)
    add_heading(doc, "形成的管理基础", 3)
    add_list(doc, [
        "业务数据进入财务体系，并支持部门与柜组层面的利润分析。",
        "返利自动化进入线上验证，为应收、ROI和绩效管理提供数据基础。",
        "关键业务规则、测试过程与管理应用形成可追踪的项目链路。",
    ])

    add_heading(doc, "5.2 品类管理试点：先建立可重复的方法", 2)
    add_kv(doc, "试点设计", "1家试点门店、2个功能分类、4轮双周复盘。")
    add_body(doc, "品类管理的难点，是把分析结论转化为门店动作，并通过复盘验证哪些方法值得保留。试点按照“分类—指标—动作—复盘”的顺序组织。")
    add_list(doc, [
        "分类：统一品类角色与消费任务。",
        "指标：定义销量、毛利和结构等观察口径。",
        "动作：形成商品调整与门店执行动作。",
        "复盘：验证结果、识别偏差并修正方法。",
    ], numbered=True)
    add_callout(doc, "实践结论", "先在小范围回答方法是否被理解、指标是否反映动作变化、流程是否具备复制条件，再决定是否扩大。")

    add_heading(doc, "5.3 购物中心IT规划：先说清总部与项目边界", 2)
    add_kv(doc, "管理命题", "哪些能力由总部统一，哪些由项目现场承担？")
    add_body(doc, "规划围绕招商、运营、物业、会员、预算和BI等能力展开，形成业务蓝图、应用架构、系统边界和建设顺序，并衔接开业筹备、培训、运营支持与项目验收。")
    add_kv(doc, "主要价值", "为管理层决策、供应商选型和分期建设提供统一框架。")

    add_heading(doc, "5.4 复杂业务产品设计：把边界与规则转化为产品", 2)
    add_kv(doc, "设计命题", "复杂业务如何被不同角色正确理解和操作？")
    add_body(doc, "围绕角色、流程、状态、权限、异常和关键数据模型开展设计，通过PRD、交互原型、业务规则和场景化验收标准形成可评审、可开发、可测试的需求与产品方案。")
    add_callout(doc, "共同原则", "文件完成并不等于项目完成。交付应同时满足管理层能决策、项目团队能执行、业务现场能验证。")


def add_method(doc):
    doc.add_page_break()
    add_heading(doc, "6. 项目方法与治理机制", 1)
    add_section_illustration(
        doc, "06-governance.png",
        "章节导读｜五步项目方法与清晰的双方责任共同控制交付质量",
        "五个项目阶段与客户、顾问两条责任路径在决策点汇合。",
    )
    add_body(doc, "咨询价值来自专业判断，也来自清晰的决策机制、责任分工和持续验证。简益通常按照以下五个环节组织项目。")
    add_list(doc, [
        "识别课题：明确项目要解决的经营或管理问题。",
        "统一事实：通过访谈、资料和数据建立共同事实基础。",
        "设计方案：形成目标模型、关键取舍与实施路径。",
        "组织实施：管理计划、范围、风险、问题和跨方协同。",
        "验证改进：通过测试、复盘、验收和现场反馈持续修正。",
    ], numbered=True)
    add_heading(doc, "6.1 双方责任分工", 2)
    add_table(doc,
              ["项目环节", "简益负责", "客户负责", "共同确认"],
              [
                  ("问题定义", "访谈、诊断与结构化表达", "提供真实资料与关键人员", "课题边界与优先级"),
                  ("方案设计", "方法、模型、流程与方案", "业务判断与约束条件", "目标方案与关键取舍"),
                  ("项目推进", "计划、风险、协调与跟踪", "资源投入与内部决策", "里程碑与问题处理"),
                  ("验证移交", "测试、复盘与移交建议", "现场执行与结果反馈", "验收标准与后续行动"),
              ],
              [1700, 2600, 2500, CONTENT_DXA - 6800], font_size=9.1)
    add_callout(doc, "项目治理重点", "把需要客户决策的问题及时提升，把需要供应商解决的问题持续跟踪，把需要现场验证的内容明确到场景和验收标准。")


def add_start_about(doc):
    add_heading(doc, "7. 如何启动合作", 1)
    add_section_illustration(
        doc, "07-start.png",
        "章节导读｜先对齐问题、边界、人员、资料与决策机制",
        "项目启动桌面整理问题、范围、参与人员、资料和决策安排。",
    )
    add_body(doc, "一次有效的项目讨论，可以从四项准备开始。资料不需要非常完整，但应尽可能真实，并由熟悉业务和系统现状的人员参与。")
    add_list(doc, [
        "列出当前最希望解决的3个问题，并说明为什么重要。",
        "明确涉及的部门、业态、门店、区域和系统范围。",
        "准备现有流程、报表、需求文档、项目计划和问题记录。",
        "确认内部负责人、决策机制以及时间、预算和资源约束。",
    ])
    add_heading(doc, "7.1 首次沟通通常要回答的问题", 2)
    add_table(doc,
              ["讨论主题", "需要确认的内容"],
              [
                  ("为什么现在做", "经营压力、管理变化、系统建设或项目风险是什么？"),
                  ("希望得到什么", "管理决策、业务改进、系统方案或项目交付的目标是什么？"),
                  ("工作范围", "哪些组织、门店、流程、数据和系统在项目范围内？"),
                  ("如何验收", "哪些交付物、业务场景和管理结果可以证明项目完成？"),
              ],
              [2200, CONTENT_DXA - 2200], font_size=9.5)

    add_heading(doc, "8. 关于简益咨询", 1)
    add_section_illustration(
        doc, "08-about.png",
        "章节导读｜连接零售现场、软件产品、项目管理与管理决策",
        "一座桥连接零售现场、产品设计空间和管理决策会议室。",
    )
    add_body(doc, "核心顾问的职业经历横跨零售企业、零售软件和新零售平台，能够同时理解业务现场、系统产品逻辑和项目交付。")
    add_table(doc,
              ["经验视角", "代表经历", "带来的理解"],
              [
                  ("零售企业", "中百集团", "理解门店、商品、运营和组织协同的真实约束"),
                  ("零售软件", "富基融通", "理解大型零售系统的产品逻辑、实施路径和交付风险"),
                  ("新零售平台", "有赞新零售", "理解数字化产品、客户场景与平台能力的连接方式"),
              ],
              [1800, 1800, CONTENT_DXA - 3600], font_size=9.4)
    add_callout(doc, "下一步", "如需讨论具体课题，请联系简益咨询商务联系人。双方可先通过一次项目访谈确认问题边界，再决定采用诊断、专项方案或实施陪跑。")
    disclaimer = doc.add_paragraph()
    disclaimer.paragraph_format.space_before = Pt(12)
    disclaimer.paragraph_format.space_after = Pt(0)
    r = disclaimer.add_run("说明：本白皮书用于业务交流。具体服务范围、人员投入、周期、价格和验收标准，以双方确认并签署的项目文件为准。")
    set_rfonts(r, size=8.5, color=MID_GRAY)


def build():
    doc = Document()
    configure_styles(doc)
    configure_section(doc)
    props = doc.core_properties
    props.title = "简益咨询产品与服务白皮书"
    props.subject = "零售经营咨询、IT规划、项目管理与产品设计"
    props.author = "简益咨询"
    props.keywords = "零售咨询, IT规划, 项目管理, 产品设计"
    props.comments = "商务文档配图版 V1.3｜中文统一使用黑体-简（Heiti SC）"

    add_cover(doc)
    start_body_section(doc)
    add_guide(doc)
    add_core(doc)
    add_service_system(doc)
    add_packages(doc)
    add_service_details(doc)
    add_cases(doc)
    add_method(doc)
    add_start_about(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
